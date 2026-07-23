"""Fixtures that build **real** PDF/DOCX/XLSX documents under a document root (mirroring
knowledge-importer's own parser tests), so the reader tools run against genuine files — no mocks."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from openpyxl import Workbook
from pipeline_contracts import TenantContext
from reportlab.pdfgen import canvas


def _make_pdf(path: Path) -> None:
    c = canvas.Canvas(str(path))
    c.drawString(72, 720, "Confidentiality Policy")
    c.showPage()
    c.drawString(72, 720, "Access control requirements")
    c.showPage()
    c.save()


def _make_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Policy Statement")
    doc.add_paragraph("This policy governs access control.")
    doc.save(str(path))


def _make_xlsx(path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    assert ws is not None
    ws.append(["Control ID", "Description"])
    ws.append(["AC-1", "Access Control Policy"])
    wb.save(str(path))


@pytest.fixture
def doc_root(tmp_path: Path) -> Path:
    """A document root holding one real file of each supported type, plus a nested sub-folder."""
    root = tmp_path / "evidence"
    (root / "policies").mkdir(parents=True)
    _make_pdf(root / "confidentiality.pdf")
    _make_docx(root / "policies" / "access-control.docx")
    _make_xlsx(root / "controls.xlsx")
    (root / "broken.pdf").write_bytes(b"not a real pdf, just garbage bytes")
    return root


@pytest.fixture
def tenant() -> TenantContext:
    return TenantContext(tenant_id="org_acme", principal_id="u1", roles=("analyst",))
