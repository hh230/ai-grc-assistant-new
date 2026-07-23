"""DOCX export — the artifact a customer downloads. We build the `.docx` bytes and re-open them with
`python-docx` to assert the real document content (not just that bytes were produced)."""

from __future__ import annotations

from datetime import datetime, timezone
from io import BytesIO

from deliverables import (
    build_deliverable,
    build_gap_matrix,
    deliverable_to_docx,
    gap_matrix_to_docx,
)
from docx import Document

_FIXED = datetime(2026, 7, 20, 12, 0, tzinfo=timezone.utc)


def _texts(data: bytes) -> list[str]:
    return [p.text for p in Document(BytesIO(data)).paragraphs]


def test_deliverable_docx_is_a_real_openable_document(gap_mission) -> None:
    data = deliverable_to_docx(build_deliverable(gap_mission, title="Gap Assessment", now=_FIXED))
    assert data[:2] == b"PK"                                 # a .docx is a zip — real file bytes
    texts = _texts(data)
    assert "Gap Assessment" in texts
    assert "Tenant: org_acme" in texts
    assert "Identify Controls" in texts
    assert any("Sources: doc-acme-1" in t for t in texts)   # provenance exported to the docx


def test_gap_matrix_docx_has_the_evidence_table(gap_mission, library) -> None:
    data = gap_matrix_to_docx(build_gap_matrix(gap_mission, library))
    document = Document(BytesIO(data))
    texts = [p.text for p in document.paragraphs]
    assert any("Gap Matrix — Evidence Mapping (ISO/IEC 27001:2022)" in t for t in texts)
    assert any("not a compliance attestation" in t for t in texts)   # honest naming in the export

    assert len(document.tables) == 1
    rows = [[c.text for c in row.cells] for row in document.tables[0].rows]
    assert rows[0] == ["Control", "Title", "Status", "Evidence"]
    assert ["A.8.5", "Secure authentication", "covered", "doc-acme-1"] in rows
    assert ["A.8.24", "Use of cryptography", "gap", "—"] in rows
