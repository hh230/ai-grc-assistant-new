"""Export deliverables to **DOCX** and **PDF** — the formats a customer downloads and shares (P3).

Each function returns the file as `bytes` (no filesystem I/O — the caller writes it wherever: an
HTTP response, an object store, a file). DOCX is built with `python-docx`, PDF with `reportlab`; both
mirror the Markdown renderer, so every export carries the same content and provenance. Pure and
deterministic given the deliverable.
"""

from __future__ import annotations

from io import BytesIO
from xml.sax.saxutils import escape

from docx import Document
from docx.document import Document as DocxDocument
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from deliverables.models import Deliverable, GapMatrix


def _to_bytes(document: DocxDocument) -> bytes:
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def deliverable_to_docx(deliverable: Deliverable) -> bytes:
    """A generic `Deliverable` as a `.docx`: a titled document with the goal/tenant/timestamp header
    and one heading + body per section, each footed by its sources (provenance) and confidence."""
    document = Document()
    document.add_heading(deliverable.title, level=0)
    document.add_paragraph(f"Goal: {deliverable.goal}")
    document.add_paragraph(f"Tenant: {deliverable.tenant_id}")
    document.add_paragraph(f"Generated: {deliverable.generated_at}")
    for section in deliverable.sections:
        document.add_heading(section.heading, level=1)
        document.add_paragraph(section.body or "(no content)")
        meta: list[str] = []
        if section.citations:
            meta.append("Sources: " + ", ".join(section.citations))
        if section.confidence is not None:
            meta.append(f"Confidence: {section.confidence:.2f}")
        if meta:
            document.add_paragraph(" · ".join(meta))
    return _to_bytes(document)


def gap_matrix_to_docx(matrix: GapMatrix) -> bytes:
    """The Gap Matrix (Evidence Mapping) as a `.docx`: the headline evidence coverage, the honesty
    note, and a `Control | Title | Status | Evidence` table — the artifact a customer downloads."""
    document = Document()
    document.add_heading(f"Gap Matrix — Evidence Mapping ({matrix.framework})", level=0)
    document.add_paragraph(f"Scope: {matrix.scope}")
    document.add_paragraph(
        f"Evidence coverage: {matrix.covered_count}/{matrix.total} "
        f"({matrix.coverage:.0%}) controls have supporting evidence in the corpus"
    )
    document.add_paragraph(f"Gaps (no evidence found): {len(matrix.gaps)}")
    document.add_paragraph(
        "Evidence mapping (lexical) — supporting evidence found, not a compliance attestation."
    )

    table = document.add_table(rows=1, cols=4)
    header = table.rows[0].cells
    header[0].text, header[1].text = "Control", "Title"
    header[2].text, header[3].text = "Status", "Evidence"
    for row in matrix.rows:
        cells = table.add_row().cells
        cells[0].text = row.control_code
        cells[1].text = row.control_title
        cells[2].text = row.status
        cells[3].text = ", ".join(row.evidence) if row.evidence else "—"
    return _to_bytes(document)


# ── PDF (reportlab) ──────────────────────────────────────────────────────────────────

def _pdf(elements: list[object]) -> bytes:
    buffer = BytesIO()
    SimpleDocTemplate(buffer, pagesize=A4, title="Rasheed deliverable").build(elements)
    return buffer.getvalue()


def deliverable_to_pdf(deliverable: Deliverable) -> bytes:
    """A generic `Deliverable` as a `.pdf`: title, goal/tenant/timestamp header, and one heading +
    body per section footed by its sources (provenance) and confidence."""
    styles = getSampleStyleSheet()
    elements: list[object] = [
        Paragraph(escape(deliverable.title), styles["Title"]),
        Paragraph(f"Goal: {escape(deliverable.goal)}", styles["Normal"]),
        Paragraph(f"Tenant: {escape(deliverable.tenant_id)}", styles["Normal"]),
        Paragraph(f"Generated: {escape(deliverable.generated_at)}", styles["Normal"]),
        Spacer(1, 12),
    ]
    for section in deliverable.sections:
        elements.append(Paragraph(escape(section.heading), styles["Heading1"]))
        elements.append(Paragraph(escape(section.body or "(no content)"), styles["Normal"]))
        meta: list[str] = []
        if section.citations:
            meta.append("Sources: " + ", ".join(section.citations))
        if section.confidence is not None:
            meta.append(f"Confidence: {section.confidence:.2f}")
        if meta:
            elements.append(Paragraph(escape(" · ".join(meta)), styles["Italic"]))
        elements.append(Spacer(1, 8))
    return _pdf(elements)


def gap_matrix_to_pdf(matrix: GapMatrix) -> bytes:
    """The Gap Matrix (Evidence Mapping) as a `.pdf`: the headline evidence coverage, the honesty
    note, and a `Control | Title | Status | Evidence` table — the artifact a customer downloads."""
    styles = getSampleStyleSheet()
    elements: list[object] = [
        Paragraph(escape(f"Gap Matrix — Evidence Mapping ({matrix.framework})"), styles["Title"]),
        Paragraph(f"Scope: {escape(matrix.scope)}", styles["Normal"]),
        Paragraph(
            f"Evidence coverage: {matrix.covered_count}/{matrix.total} "
            f"({matrix.coverage:.0%}) controls have supporting evidence in the corpus",
            styles["Normal"],
        ),
        Paragraph(f"Gaps (no evidence found): {len(matrix.gaps)}", styles["Normal"]),
        Paragraph(
            "Evidence mapping (lexical) — supporting evidence found, not a compliance attestation.",
            styles["Italic"],
        ),
        Spacer(1, 12),
    ]
    data = [["Control", "Title", "Status", "Evidence"]]
    for row in matrix.rows:
        evidence = ", ".join(row.evidence) if row.evidence else "—"
        data.append([row.control_code, row.control_title, row.status, evidence])
    table = Table(data, repeatRows=1)
    table.setStyle(
        TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e6e6e6")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ])
    )
    elements.append(table)
    return _pdf(elements)
