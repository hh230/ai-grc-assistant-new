from __future__ import annotations

from pathlib import Path

from docx import Document
from openpyxl import Workbook
from reportlab.pdfgen import canvas

from knowledge_importer.cli import build_pipeline
from knowledge_importer.models import document_id_for


def _make_pdf(path: Path) -> None:
    c = canvas.Canvas(str(path))
    c.drawString(72, 720, "Confidentiality Policy")
    c.showPage()
    c.drawString(72, 720, "Page two content")
    c.showPage()
    c.save()


def _make_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Policy Statement")
    doc.add_paragraph("This policy governs access control.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Control"
    table.rows[0].cells[1].text = "Description"
    doc.save(str(path))


def _make_xlsx(path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    assert ws is not None
    ws.title = "Sheet1"
    ws.append(["Control ID", "Description"])
    ws.append(["AC-1", "Access Control Policy"])
    wb.create_sheet("Sheet2")
    wb["Sheet2"].append(["Extra", "Data"])
    wb.save(str(path))


def _build_mixed_library(root: Path) -> None:
    (root / "policies").mkdir(parents=True)
    (root / "policies" / "notes.txt").write_text("Plain text notes on retention.")
    (root / "policies" / "readme.md").write_text("# Overview\n\nMarkdown body.")
    _make_pdf(root / "policies" / "confidentiality.pdf")
    _make_docx(root / "policies" / "access-control.docx")
    _make_xlsx(root / "policies" / "controls.xlsx")
    (root / "policies" / "broken.pdf").write_bytes(b"not a real pdf, just garbage bytes")


def test_every_supported_format_parses_successfully(tmp_path: Path) -> None:
    library_root = tmp_path / "library"
    imports_dir = tmp_path / "imports"
    _build_mixed_library(library_root)

    run = build_pipeline(imports_dir=imports_dir, chunks_dir=tmp_path / "chunks").run(library_root)
    by_relpath = {m.relative_path: m for m in run.manifests}

    txt = by_relpath["policies/notes.txt"]
    assert txt.parsed is True
    assert txt.parser == "text"
    assert txt.page_count is None
    assert txt.character_count == len("Plain text notes on retention.")

    md = by_relpath["policies/readme.md"]
    assert md.parsed is True
    assert md.parser == "markdown"
    assert md.page_count is None
    assert "Markdown body." in (imports_dir / f"{md.document_id}.txt").read_text(encoding="utf-8")

    pdf = by_relpath["policies/confidentiality.pdf"]
    assert pdf.parsed is True
    assert pdf.parser == "pdf"
    assert pdf.page_count == 2
    extracted_pdf_text = (imports_dir / f"{pdf.document_id}.txt").read_text(encoding="utf-8")
    assert "Confidentiality Policy" in extracted_pdf_text
    assert "Page two content" in extracted_pdf_text

    docx = by_relpath["policies/access-control.docx"]
    assert docx.parsed is True
    assert docx.parser == "docx"
    assert docx.page_count is None
    extracted_docx_text = (imports_dir / f"{docx.document_id}.txt").read_text(encoding="utf-8")
    assert "Policy Statement" in extracted_docx_text
    assert "Control" in extracted_docx_text and "Description" in extracted_docx_text

    xlsx = by_relpath["policies/controls.xlsx"]
    assert xlsx.parsed is True
    assert xlsx.parser == "xlsx"
    assert xlsx.page_count == 2
    extracted_xlsx_text = (imports_dir / f"{xlsx.document_id}.txt").read_text(encoding="utf-8")
    assert "AC-1" in extracted_xlsx_text
    assert "Extra" in extracted_xlsx_text


def test_all_parsed_manifests_record_required_fields(tmp_path: Path) -> None:
    library_root = tmp_path / "library"
    imports_dir = tmp_path / "imports"
    _build_mixed_library(library_root)

    run = build_pipeline(imports_dir=imports_dir, chunks_dir=tmp_path / "chunks").run(library_root)
    successful = [m for m in run.manifests if m.relative_path != "policies/broken.pdf"]

    for manifest in successful:
        assert manifest.parsed is True
        assert manifest.parser is not None
        assert manifest.parser_used is not None
        assert manifest.parser_fallback is False  # clean fixtures parse on the primary backend
        assert manifest.parser_attempts and manifest.parser_attempts[-1]["ok"] is True
        assert manifest.failure_reason is None
        assert manifest.character_count is not None and manifest.character_count > 0
        assert manifest.extraction_duration is not None and manifest.extraction_duration >= 0
        assert manifest.parsed_at
        assert manifest.error is None
        assert manifest.status == "parsed"
        assert manifest.stages_completed == ("intake", "parsing", "profile_assignment", "chunking")
        extracted = imports_dir / f"{manifest.document_id}.txt"
        assert extracted.exists()

    # the clean PDF fixture is produced by reportlab, which pypdf reads fine -> primary backend
    pdf = next(m for m in successful if m.relative_path == "policies/confidentiality.pdf")
    assert pdf.parser == "pdf"
    assert pdf.parser_used == "pypdf"


def test_parse_failure_is_recorded_and_does_not_stop_the_pipeline(tmp_path: Path) -> None:
    library_root = tmp_path / "library"
    imports_dir = tmp_path / "imports"
    _build_mixed_library(library_root)

    run = build_pipeline(imports_dir=imports_dir, chunks_dir=tmp_path / "chunks").run(library_root)
    by_relpath = {m.relative_path: m for m in run.manifests}

    broken = by_relpath["policies/broken.pdf"]
    assert broken.parsed is False
    assert broken.error is not None and broken.error != ""
    assert broken.status == "parse_failed"
    assert broken.parser == "pdf"
    assert broken.failure_reason is not None
    # a genuinely-invalid PDF exhausts BOTH backends before failing
    attempted = [a["backend"] for a in broken.parser_attempts]
    assert attempted == ["pypdf", "pypdfium2"]
    assert all(a["ok"] is False for a in broken.parser_attempts)
    assert not (imports_dir / f"{broken.document_id}.txt").exists()

    # every other document in the same run still parsed successfully
    others = [m for path, m in by_relpath.items() if path != "policies/broken.pdf"]
    assert len(others) == 5
    assert all(m.parsed for m in others)


def test_rerun_reparses_with_stable_document_ids(tmp_path: Path) -> None:
    library_root = tmp_path / "library"
    imports_dir = tmp_path / "imports"
    _build_mixed_library(library_root)

    pipeline = build_pipeline(imports_dir=imports_dir, chunks_dir=tmp_path / "chunks")
    first_run = pipeline.run(library_root)
    second_run = pipeline.run(library_root)

    first_by_id = {m.document_id: m for m in first_run.manifests}
    second_by_id = {m.document_id: m for m in second_run.manifests}
    assert first_by_id.keys() == second_by_id.keys()
    for document_id in first_by_id:
        assert document_id == document_id_for(first_by_id[document_id].relative_path)
        assert first_by_id[document_id].parsed == second_by_id[document_id].parsed
